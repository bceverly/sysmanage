#!/usr/bin/env python3
"""Build the SysManage Executive Opportunity Brief as a professional .docx.

Reads   marketing/executive-opportunity-brief.md
Writes  marketing/executive-opportunity-brief.docx

Run it again any time the Markdown changes:

    .venv/bin/python scripts/build-marketing-brief.py

Why this exists instead of a one-line pandoc call:
  * We want branded, layout-controlled output (cover page, colored headings,
    callout boxes, styled tables, footer with page numbers).
  * SVG figures are rasterized cleanly via  SVG -> PDF (LibreOffice) -> PNG
    (Ghostscript @300dpi).  ImageMagick's built-in SVG renderer produces the
    white horizontal seam lines we keep fighting; this pipeline does not.

External tools required (all already present on this machine):
    libreoffice / soffice   - SVG -> PDF
    gs (Ghostscript)        - PDF -> PNG
Python deps: python-docx, Pillow  (in the project .venv)
"""

import os
import re
import shutil
import subprocess
import sys
import tempfile

from PIL import Image

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Paths & brand
# --------------------------------------------------------------------------- #
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MK = os.path.join(REPO, "marketing")
MD_PATH = os.path.join(MK, "executive-opportunity-brief.md")
DOCX_PATH = os.path.join(MK, "executive-opportunity-brief.docx")
PDF_PATH = os.path.join(MK, "executive-opportunity-brief.pdf")
BUILD = os.path.join(MK, ".build")  # rasterized SVGs land here (gitignore-able)

# Brand palette (from sysmanage-logo.svg / product diagrams)
BRAND = RGBColor(0x19, 0x76, 0xD2)      # primary blue
ACCENT = RGBColor(0x38, 0x8E, 0x3C)     # agent green
DARK = RGBColor(0x2C, 0x3E, 0x50)       # heading slate
BODY = RGBColor(0x33, 0x33, 0x33)       # body text
MUTED = RGBColor(0x6C, 0x75, 0x7D)      # captions / footer
BRAND_HEX = "1976D2"
HEADER_FILL = "1976D2"
CALLOUT_FILL = "EAF3FB"
ZEBRA_FILL = "F5F8FB"

CONTENT_WIDTH_IN = 6.6   # usable width between margins
MAX_IMG_IN = 6.5

# --------------------------------------------------------------------------- #
# Low-level OOXML helpers
# --------------------------------------------------------------------------- #

def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _cell_borders(cell, **edges):
    """edges: e.g. left={'sz':24,'color':'1976D2'} ; sz in eighths of a point."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        e = OxmlElement(f"w:{edge}")
        if spec is None:
            e.set(qn("w:val"), "nil")
        else:
            e.set(qn("w:val"), spec.get("val", "single"))
            e.set(qn("w:sz"), str(spec.get("sz", 4)))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), spec.get("color", "auto"))
        borders.append(e)
    tcPr.append(borders)


def _bottom_rule(paragraph, color=BRAND_HEX, sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)


def _no_table_autofit_full_width(table):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _add_page_number(paragraph):
    def fld(t):
        e = OxmlElement("w:fldChar")
        e.set(qn("w:fldCharType"), t)
        return e
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(fld("begin"))
    run._r.append(instr)
    run._r.append(fld("end"))


# --------------------------------------------------------------------------- #
# Third-party trademark marking
# --------------------------------------------------------------------------- #
# Every occurrence of a third-party company/product name is automatically tagged
# with the correct symbol at build time, so the Markdown source stays clean and
# editable.  ® = marks widely known to be registered; ™ = product sub-names,
# projects, and firms where registration isn't asserted here.  A global
# "Trademarks & Legal Notice" section in the Markdown carries the attributions.
# NOTE: names baked into the rasterized SVG diagrams (e.g. PostgreSQL, Windows in
# the architecture image) are pixels, not text — they are covered by the notice.
_R, _T = "®", "™"
TRADEMARKS = {
    # Microsoft
    "Microsoft": _R, "Windows": _R, "Intune": _R,
    "Configuration Manager": _T, "MECM": _T, "SCCM": _T,
    # Apple
    "Apple": _R, "macOS": _R,
    # Jamf
    "Jamf": _R,
    # Red Hat
    "Red Hat Enterprise Linux": _R, "Red Hat Satellite": _R, "Red Hat": _R,
    "RHEL": _R, "Ansible": _R,
    # Canonical  (only the full "Canonical Landscape" — bare "Landscape" is generic
    # English, e.g. the "Competitive Landscape" section heading)
    "Canonical Landscape": _T, "Canonical": _R, "Ubuntu": _R,
    # SUSE
    "SUSE Linux Enterprise": _R, "SUSE Manager": _T, "SUSE": _R,
    # Operating systems / open-source components
    "Linux": _R, "FreeBSD": _R, "NetBSD": _R, "OpenBSD": _T,
    "PostgreSQL": _R, "OpenBAO": _T,
    # Security / endpoint / RMM vendors
    "Tanium": _R, "BigFix": _R, "NinjaOne": _R, "ConnectWise": _R,
    "Datto RMM": _R, "Datto": _R, "Automox": _T, "Action1": _T,
    "Qualys": _R, "Rapid7": _R, "Tenable": _R,
    # Market-research firms
    "Grand View Research": _T, "Fortune Business Insights": _T,
    "MarketsandMarkets": _T, "Mordor Intelligence": _T,
    "Virtue Market Research": _T, "Polaris Market Research": _T, "Polaris": _T,
    "Research and Markets": _T, "Custom Market Insights": _T,
}
# Longest names first so e.g. "Red Hat Enterprise Linux" wins over "Red Hat"/"Linux".
_TM_RE = re.compile(
    r"(?<!\w)(" + "|".join(re.escape(k) for k in sorted(TRADEMARKS, key=len, reverse=True)) + r")(?!\w)"
)


def mark_trademarks(text):
    return _TM_RE.sub(lambda m: m.group(1) + TRADEMARKS[m.group(1)], text)


# --------------------------------------------------------------------------- #
# Inline markdown -> runs  (**bold**, *italic*, `code`, [text](url), <autolink>)
# --------------------------------------------------------------------------- #
_INLINE = re.compile(
    r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+?\]\([^)]+?\)|<https?://[^>\s]+>)"
)


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BRAND_HEX)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(color)
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_runs(paragraph, text, size=None, color=None, bold=None):
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = paragraph.add_run(mark_trademarks(piece[2:-2]))
            r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])  # code spans left verbatim
            r.font.name = "Consolas"
        elif piece.startswith("[") and "](" in piece:
            label, url = re.match(r"\[(.+?)\]\((.+?)\)", piece).groups()
            _add_hyperlink(paragraph, url, label)
            continue
        elif piece.startswith("<") and piece.endswith(">") and "://" in piece:
            url = piece[1:-1]
            _add_hyperlink(paragraph, url, url)
            continue
        elif piece.startswith("*") and piece.endswith("*"):
            r = paragraph.add_run(mark_trademarks(piece[1:-1]))
            r.italic = True
        else:
            r = paragraph.add_run(mark_trademarks(piece))
        if size is not None:
            r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        if bold is not None and r.bold is None:
            r.bold = bold


# --------------------------------------------------------------------------- #
# SVG -> PNG (clean: SVG -> PDF via LibreOffice -> PNG via Ghostscript @300dpi)
# --------------------------------------------------------------------------- #
_SOFFICE = shutil.which("soffice") or shutil.which("libreoffice")
_GS = shutil.which("gs")


def rasterize_svg(svg_path):
    os.makedirs(BUILD, exist_ok=True)
    base = os.path.splitext(os.path.basename(svg_path))[0]
    png_path = os.path.join(BUILD, base + ".png")
    if os.path.exists(png_path) and os.path.getmtime(png_path) >= os.path.getmtime(svg_path):
        return png_path
    if not _SOFFICE or not _GS:
        raise RuntimeError("Need both libreoffice/soffice and gs to rasterize SVGs.")
    with tempfile.TemporaryDirectory() as tmp:
        profile = "file://" + os.path.join(tmp, "loprofile")
        subprocess.run(
            [_SOFFICE, "--headless", f"-env:UserInstallation={profile}",
             "--convert-to", "pdf", "--outdir", tmp, svg_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        pdf_path = os.path.join(tmp, base + ".pdf")
        subprocess.run(
            [_GS, "-dSAFER", "-dBATCH", "-dNOPAUSE", "-sDEVICE=png16m", "-r300",
             "-dTextAlphaBits=4", "-dGraphicsAlphaBits=4",
             f"-sOutputFile={png_path}", pdf_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
    return png_path


def export_pdf(docx_path):
    """Convert the finished .docx to PDF via LibreOffice (headless)."""
    if not _SOFFICE:
        raise RuntimeError("Need libreoffice/soffice to export PDF.")
    with tempfile.TemporaryDirectory() as tmp:
        profile = "file://" + os.path.join(tmp, "loprofile")
        subprocess.run(
            [_SOFFICE, "--headless", f"-env:UserInstallation={profile}",
             "--convert-to", "pdf", "--outdir", os.path.dirname(docx_path), docx_path],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )


def resolve_image(path_in_md):
    src = os.path.normpath(os.path.join(MK, path_in_md))
    if src.lower().endswith(".svg"):
        return rasterize_svg(src)
    return src


def image_width_in(png_path, cap=MAX_IMG_IN):
    with Image.open(png_path) as im:
        w, h = im.size
    natural = w / 96.0
    return min(cap, natural)


# --------------------------------------------------------------------------- #
# Block builders
# --------------------------------------------------------------------------- #

def add_image(doc, path_in_md, width_in=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    png = resolve_image(path_in_md)
    if width_in is None:
        width_in = image_width_in(png)
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(png, width=Inches(width_in))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    add_runs(p, text.strip("*"), size=8.5, color=MUTED)
    for r in p.runs:
        r.italic = True


def add_callout(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    _no_table_autofit_full_width(tbl)
    cell = tbl.cell(0, 0)
    _shade(cell, CALLOUT_FILL)
    _set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    _cell_borders(
        cell,
        left={"sz": 30, "color": BRAND_HEX},
        top={"sz": 4, "color": "D6E6F5"},
        bottom={"sz": 4, "color": "D6E6F5"},
        right={"sz": 4, "color": "D6E6F5"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_runs(p, text, size=11.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_md_table(doc, header, aligns, rows):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_autofit_full_width(tbl)
    amap = {"c": WD_ALIGN_PARAGRAPH.CENTER, "r": WD_ALIGN_PARAGRAPH.RIGHT,
            "l": WD_ALIGN_PARAGRAPH.LEFT}

    hdr = tbl.rows[0].cells
    for i, htext in enumerate(header):
        _shade(hdr[i], HEADER_FILL)
        _set_cell_margins(hdr[i])
        _cell_borders(hdr[i], bottom={"sz": 8, "color": BRAND_HEX})
        p = hdr[i].paragraphs[0]
        p.alignment = amap.get(aligns[i], WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_after = Pt(0)
        add_runs(p, htext, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        for r in p.runs:
            r.bold = True

    for ridx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, ctext in enumerate(row):
            if ridx % 2 == 1:
                _shade(cells[i], ZEBRA_FILL)
            _set_cell_margins(cells[i])
            _cell_borders(cells[i], bottom={"sz": 4, "color": "DDE3E8"})
            p = cells[i].paragraphs[0]
            p.alignment = amap.get(aligns[i], WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            add_runs(p, ctext, size=10, color=BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    _keep_with_next(p)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        add_runs(p, text, size=16, color=BRAND)
        for r in p.runs:
            r.bold = True
        _bottom_rule(p, BRAND_HEX, sz=6)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, text, size=12.5, color=DARK)
        for r in p.runs:
            r.bold = True


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    add_runs(p, text, size=11, color=BODY)


def add_list_item(doc, text, ordered):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text, size=11, color=BODY)


# --------------------------------------------------------------------------- #
# Markdown parsing
# --------------------------------------------------------------------------- #

def strip_comments(lines):
    """Remove HTML comments. A comment whose entire content is exactly
    ``PAGEBREAK`` (i.e. ``<!-- PAGEBREAK -->``) becomes a page-break sentinel.
    Comments that merely mention the word (e.g. the doc-convention notes at the
    top of the file) are ignored."""
    out = []
    in_comment = False
    for line in lines:
        s = line.rstrip("\n")
        if in_comment:
            # Page-break markers are always single-line, so multi-line comment
            # bodies are simply skipped.
            if "-->" in s:
                s = s[s.index("-->") + 3:]
                in_comment = False
            else:
                continue
        # collapse all fully-closed comments on this line
        while "<!--" in s and "-->" in s:
            a = s.index("<!--")
            b = s.index("-->", a) + 3
            inner = s[a + 4:b - 3].strip()
            if inner == "PAGEBREAK":
                out.append("\x00PAGEBREAK")
            s = s[:a] + s[b:]
        if "<!--" in s:  # a comment opens and does not close on this line
            s = s[:s.index("<!--")]
            in_comment = True
        out.append(s)
    return out


def is_table_sep(line):
    return "---" in line and re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line) is not None


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def parse_aligns(sep_cells):
    out = []
    for c in sep_cells:
        c = c.strip()
        if c.startswith(":") and c.endswith(":"):
            out.append("c")
        elif c.endswith(":"):
            out.append("r")
        else:
            out.append("l")
    return out


IMG_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def render_cover(doc, lines):
    for _ in range(3):
        doc.add_paragraph()
    first_note = True
    for s in lines:
        if not s.strip():
            continue
        m = IMG_RE.match(s.strip())
        if m:
            path = m.group(2)
            if "logo" in path.lower():
                add_image(doc, path, width_in=2.0)
            else:
                add_image(doc, path, width_in=5.6)
            continue
        if s.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, s[2:], size=32, color=BRAND)
            for r in p.runs:
                r.bold = True
            continue
        if s.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(14)
            add_runs(p, s[3:], size=15, color=DARK)
            continue
        if s.startswith("**") and s.endswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(s.strip("*").upper())
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = MUTED
            continue
        if s.startswith("*") and s.endswith("*"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24) if first_note else Pt(4)
            first_note = False
            r = p.add_run(s.strip("*"))
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = MUTED
            continue
    doc.add_page_break()


def render_body(doc, lines):
    i = 0
    n = len(lines)
    para = []

    def flush():
        nonlocal para
        if para:
            add_body_paragraph(doc, " ".join(para))
            para = []

    while i < n:
        s = lines[i]
        stripped = s.strip()

        if s == "\x00PAGEBREAK":
            flush()
            doc.add_page_break()
            i += 1
            continue
        if not stripped:
            flush()
            i += 1
            continue
        # table?
        if "|" in s and i + 1 < n and is_table_sep(lines[i + 1]):
            flush()
            header = split_row(s)
            aligns = parse_aligns(split_row(lines[i + 1]))
            rows = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                rows.append(split_row(lines[j]))
                j += 1
            add_md_table(doc, header, aligns, rows)
            i = j
            continue
        # image
        m = IMG_RE.match(stripped)
        if m:
            flush()
            add_image(doc, m.group(2))
            i += 1
            continue
        # headings
        if s.startswith("### "):
            flush()
            add_heading(doc, s[4:], 2)
            i += 1
            continue
        if s.startswith("## "):
            flush()
            add_heading(doc, s[3:], 1)
            i += 1
            continue
        if s.startswith("# "):
            flush()
            add_heading(doc, s[2:], 1)
            i += 1
            continue
        # blockquote / callout (collect consecutive)
        if s.startswith(">"):
            flush()
            buf = []
            while i < n and lines[i].startswith(">"):
                buf.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            add_callout(doc, " ".join(x.strip() for x in buf if x.strip()))
            continue
        # list items
        if re.match(r"^-\s+", s):
            flush()
            add_list_item(doc, re.sub(r"^-\s+", "", s), ordered=False)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", s):
            flush()
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", s), ordered=True)
            i += 1
            continue
        # horizontal rule (non-table)
        if stripped in ("---", "***", "___"):
            flush()
            i += 1
            continue
        # small italic note line
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            flush()
            add_caption(doc, stripped)
            i += 1
            continue
        # ordinary text
        para.append(stripped)
        i += 1
    flush()


# --------------------------------------------------------------------------- #
# Document setup
# --------------------------------------------------------------------------- #

def setup_document():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY

    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)

    # Footer line 1: "SysManage · Executive Opportunity Brief" (left) | page # (right)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    tab_pos = Inches(CONTENT_WIDTH_IN)
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(tab_pos.twips)))
    tabs.append(tab)
    pPr.append(tabs)
    r = fp.add_run("SysManage  ·  Executive Opportunity Brief")
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    fp.add_run("\t")
    r2 = fp.add_run("")
    r2.font.size = Pt(8)
    r2.font.color.rgb = MUTED
    _add_page_number(fp)

    # Footer line 2: centered confidentiality + copyright notice
    cp = footer.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cr = cp.add_run(
        "Confidential  ·  © 2026 Bryan Everly. All rights reserved."
    )
    cr.font.size = Pt(7.5)
    cr.font.color.rgb = MUTED
    cr.italic = True
    return doc


def main():
    for tool, name in ((_SOFFICE, "libreoffice/soffice"), (_GS, "gs (ghostscript)")):
        if not tool:
            sys.exit(f"ERROR: required tool not found: {name}")
    if not os.path.exists(MD_PATH):
        sys.exit(f"ERROR: source not found: {MD_PATH}")

    with open(MD_PATH, encoding="utf-8") as f:
        raw = f.readlines()
    lines = strip_comments(raw)

    # Split cover (everything before first PAGEBREAK) from body.
    try:
        pb = lines.index("\x00PAGEBREAK")
        cover_lines, body_lines = lines[:pb], lines[pb + 1:]
    except ValueError:
        cover_lines, body_lines = [], lines

    doc = setup_document()
    render_cover(doc, cover_lines)
    render_body(doc, body_lines)
    doc.save(DOCX_PATH)
    print(f"Wrote {os.path.relpath(DOCX_PATH, REPO)}")

    export_pdf(DOCX_PATH)
    print(f"Wrote {os.path.relpath(PDF_PATH, REPO)}")
    print(f"Rasterized SVGs cached in {os.path.relpath(BUILD, REPO)}/")


if __name__ == "__main__":
    main()
